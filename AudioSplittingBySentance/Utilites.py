from pydub import AudioSegment, silence
from pydub.silence import split_on_silence
from docx import Document
import pydub

import numpy as np
import librosa
from dtw import dtw
import wave
import librosa

import scipy.spatial.distance
from sklearn.feature_extraction.text import TfidfVectorizer

import Config


def get_pauses(audio_file_path):
    myaudio = AudioSegment.from_wav(audio_file_path)
    # get the average loudness of the audio
    dBFS = myaudio.dBFS
    silence = pydub.silence.detect_silence(myaudio, min_silence_len=500, silence_thresh=dBFS - 16)
    return [(stop - start) for start, stop in silence]


def get_average_loudness(audio_file_path):
    audio = AudioSegment.from_wav(audio_file_path)
    return audio.dBFS - 16

def split_sentences(text):
    sentences = text.split('. ')
    sentences_with_colon = []
    for sentence in sentences:
        if Config.split_character in sentence:
            sentences_with_colon.extend([s.strip() + Config.split_character for s in sentence.split(Config.split_character)])
        else:
            sentences_with_colon.append(sentence.strip())
    return sentences_with_colon


def read_docx_file(filepath):
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]
    text = ' '.join(paragraphs)
    return text


def get_audio_speaker_speed(audio_sentence, sample_rate):
    # Calculate the duration of the audio sentence in seconds
    audio_duration = len(audio_sentence) / sample_rate
    # Calculate the number of words in the audio sentence by counting the spaces
    audio_words = audio_sentence.count(" ") + 1
    # Calculate the speed of the speaker in words per minute
    audio_speed = audio_words / audio_duration * 60
    return audio_speed


def get_text_speed(text, wpm=150):
    # Calculate the number of words in the text sentence by counting the spaces
    text_words = text.count(" ") + 1

    # Calculate the speed of the text sentence in words per minute, assuming a standard speed of 150 wpm
    text_speed = text_words / wpm * 60

    return text_speed


def get_audio_chunks(audio_path, silence_len, thresh):
    # reading from audio mp3 file
    sound = AudioSegment.from_mp3(audio_path)
    # splitting audio files
    audio_chunks = split_on_silence(sound, min_silence_len=silence_len, silence_thresh=thresh)
    return audio_chunks


def get_audio_sample_rate(audio_path):
    sound = AudioSegment.from_mp3(audio_path)
    return sound.frame_rate


def save_audio_files(audio_chunks, output_dir):
    # loop is used to iterate over the output list
    saved_audio_file_paths = []
    for i, chunk in enumerate(audio_chunks):
        out_file = output_dir + "segment_{}.wav".format(i)
        # print("Exporting file", out_file)
        chunk.export(out_file, format="wav")
        saved_audio_file_paths.append(out_file)
    return saved_audio_file_paths


def save_audio_file(audio, file_path):
    audio.export(file_path, format="wav")


def convert_to_wav(mp3_file, wav_file):
    audio = pydub.AudioSegment.from_mp3(mp3_file)
    audio.export(wav_file, format='wav')


def get_WPS(audio_file_path, text_file_path):
    # Extract text from DOCX file
    document = Document(text_file_path)
    text = ' '.join([paragraph.text for paragraph in document.paragraphs])

    # Load the audio file and obtain its duration
    with wave.open(audio_file_path, 'rb') as wav:
        frame_rate = wav.getframerate()
        n_frames = wav.getnframes()
        audio_duration = n_frames / frame_rate

    # Count the number of words in the text
    word_count = len(text.split())

    # Compute the average words per second
    words_per_second = word_count / audio_duration

    return words_per_second

def feature_based_matching(audio_file_path, text_sentence):
    # Extract the MFCC features from the audio
    audio, sample_rate = librosa.load(audio_file_path)
    mfcc = librosa.feature.mfcc(y=audio, sr=sample_rate)

    # Extract the TF-IDF features from the text
    vectorizer = TfidfVectorizer()
    tfidf = vectorizer.fit_transform([text_sentence]).toarray()

    # Align the MFCC matrix and the TF-IDF vector using DTW
    dist, cost, path = dtw(mfcc.T, tfidf)
    mfcc_aligned = mfcc[:, path[0]]
    tfidf_aligned = tfidf[0, path[1]]

    # Compute the cosine similarity between the aligned features
    similarity = 1 - scipy.spatial.distance.cosine(mfcc_aligned.flatten(), tfidf_aligned.flatten())
    return similarity

def calculate_audio_duration(audio_file):
    with wave.open(audio_file, 'rb') as wav:
        frame_rate = wav.getframerate()
        n_frames = wav.getnframes()
        audio_duration = n_frames / frame_rate

    return audio_duration

def calculate_expected_duration(text, words_per_second):
    word_count = len(text.split())
    expected_duration = word_count / words_per_second

    return expected_duration

def compare_text_audio(text, audio_file, words_per_second, tolerance=1.0):
    audio_duration = calculate_audio_duration(audio_file)
    print(f"Audio duartion for: {audio_file} is: {audio_duration}")
    expected_duration = calculate_expected_duration(text, words_per_second)
    print(f"Expected duartion for: {text} is: {expected_duration}")

    duration_difference = abs(audio_duration - expected_duration)
    match = duration_difference <= tolerance

    return match