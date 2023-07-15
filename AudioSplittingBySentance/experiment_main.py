import pydub
from pydub.silence import detect_nonsilent
from difflib import SequenceMatcher
from docx import Document

import Config


def detect_pause(audio_file, min_silence_duration=500, silence_threshold=-50):
    audio = pydub.AudioSegment.from_file(audio_file)
    non_silent_regions = detect_nonsilent(audio, min_silence_duration, silence_threshold)

    pauses = []
    prev_end = 0
    for (start, end) in non_silent_regions:
        if start > prev_end:
            pauses.append((prev_end, start))
        prev_end = end

    return pauses

def calculate_speaker_speed(audio_file, transcription):
    audio = pydub.AudioSegment.from_file(audio_file)
    audio_duration = len(audio) / 1000  # duration in seconds

    document = Document(transcription)
    text = ' '.join([paragraph.text for paragraph in document.paragraphs])
    text_duration = len(text.split()) / 3  # assuming an average reading speed of 3 words per second

    speaker_speed = text_duration / audio_duration
    return speaker_speed

def match_sentences(audio_file, transcription):
    pauses = detect_pause(audio_file)
    speaker_speed = calculate_speaker_speed(audio_file, transcription)

    document = Document(transcription)
    text = ' '.join([paragraph.text for paragraph in document.paragraphs])
    text_sentences = text.splitlines()
    audio = pydub.AudioSegment.from_file(audio_file)

    audio_sentences = []
    matched_text_sentences = []

    prev_end = 0
    for pause_start, pause_end in pauses:
        audio_sentence = audio[prev_end:pause_end]
        audio_sentences.append(audio_sentence)

        text_sentence = None
        best_ratio = 0

        for sentence in text_sentences:
            ratio = SequenceMatcher(None, str(audio_sentence), sentence).ratio()
            if ratio > best_ratio:
                text_sentence = sentence
                best_ratio = ratio

        if text_sentence:
            text_sentences.remove(text_sentence)
        else:
            # If no matching sentence found based on text, estimate position using speaker speed
            estimated_duration = len(audio_sentence) / (1000 * speaker_speed)
            estimated_sentence_index = round(estimated_duration * len(text) / audio.duration_seconds)
            text_sentence = text_sentences[estimated_sentence_index]

        matched_text_sentences.append(text_sentence)
        prev_end = pause_end

    return audio_sentences, matched_text_sentences

# Usage example
audio_file = 'audio.wav'
transcription = 'transcription.docx'

audio_sentences, text_sentences = match_sentences(Config.audio_file_path, Config.docx_file_path)

for audio_sentence, text_sentence in zip(audio_sentences, text_sentences):
    print(f"Audio Sentence: {audio_sentence}")
    print(f"Text Sentence: {text_sentence}")
    print("-" * 50)
