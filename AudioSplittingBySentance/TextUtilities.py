from docx import Document
import wave
import Config
import logging

def split_sentences(text):
    sentences = text.split('. ')
    sentences_with_colon = []
    for sentence in sentences:
        if Config.split_character in sentence:
            sentences_with_colon.extend([s.strip() + Config.split_character for s in sentence.split(Config.split_character)])
        else:
            sentences_with_colon.append(sentence.strip())
    return sentences_with_colon


def read_docx_file(filepath):
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs]
    text = ' '.join(paragraphs)
    return text


def get_text_speed(text, wpm=150):
    text_words = text.count(" ") + 1

    # Calculate the speed of the text sentence in words per minute, assuming a standard speed of 150 wpm
    text_speed = text_words / wpm * 60

    return text_speed


def get_WPS(audio_file_path, text_file_path):
    # Extract text from DOCX file
    document = Document(text_file_path)
    text = ' '.join([paragraph.text for paragraph in document.paragraphs])

    # Load the audio file and obtain its duration
    with wave.open(audio_file_path, 'rb') as wav:
        frame_rate = wav.getframerate()
        n_frames = wav.getnframes()
        audio_duration = n_frames / frame_rate

    # Count the number of words in the text
    word_count = len(text.split())

    # Compute the average words per second
    words_per_second = word_count / audio_duration

    return words_per_second


def calculate_expected_duration(text, words_per_second):
    word_count = len(text.split())
    expected_duration = word_count / words_per_second

    return expected_duration

def pauses_expectant(text):
     comma_count = text.count(',')
     but_count = text.count('՝')
     ev_count = text.count(" և ")
     u_count = text.count(" ու ")
     gic_count = text.count("—")

     pauses = comma_count+but_count+ev_count+u_count+gic_count
     logging.warning("-------------------------------")
     logging.warning(f"text: {text}")
     logging.warning(f"comma_count: {comma_count}")
     logging.warning(f"but_count: {but_count}")
     logging.warning(f"ev_count: {ev_count}")
     logging.warning(f"u_count: {u_count}")
     logging.warning(f"gic_count: {gic_count}")
     logging.warning(f"expected pauses count: {pauses}")
     logging.warning("-------------------------------")

     return comma_count + but_count + ev_count + u_count + gic_count